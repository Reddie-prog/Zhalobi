"""Генерация расширенного отчёта по проекту ЖКХ.Онлайн"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Поля страницы ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(1.5)

# ── Утилиты ────────────────────────────────────────────────────────────────
BLUE      = RGBColor(37,  99,  235)
DARK      = RGBColor(15,  23,  42)
SLATE     = RGBColor(30,  41,  59)
GRAY      = RGBColor(71,  85, 105)
RED       = RGBColor(185,  28,  28)
GREEN     = RGBColor(21, 128,  61)

def tnr(run, size=11, bold=False, color=None, italic=False):
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    tnr(run, 14, bold=True, color=BLUE)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    tnr(run, 12, bold=True, color=SLATE)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    tnr(run, 11, bold=True, color=GRAY)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    return p

def para(text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    tnr(run, size, bold, color, italic)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f'{"    " * level}• {text}')
    tnr(run, size)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5 * (level + 1))
    return p

def set_cell(cell, text, bold=False, bg=None, size=10,
             align=WD_ALIGN_PARAGRAPH.LEFT, color=None, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    tnr(run, size, bold, color, italic)
    if bg:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:fill'), bg)
        shd.set(qn('w:val'),  'clear')
        tcPr.append(shd)

def table_hdr(tbl, headers, bg='1E3A5F'):
    row = tbl.rows[0]
    for j, h in enumerate(headers):
        set_cell(row.cells[j], h, bold=True, bg=bg, size=10,
                 color=RGBColor(255,255,255))

def role_chip(text):
    p = doc.add_paragraph()
    run = p.add_run(f'[ Роли: {text} ]')
    tnr(run, 9, bold=True, color=BLUE, italic=True)
    p.paragraph_format.space_after = Pt(1)
    return p

def divider():
    p = doc.add_paragraph()
    run = p.add_run(chr(9472) * 90)
    tnr(run, 8, color=RGBColor(203,213,225))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

# ══════════════════════════════════════════════════════════════════════════
#  ТИТУЛЬНАЯ СТРАНИЦА
# ══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ОТЧЁТ О ПРОЕКТНОЙ РАБОТЕ')
tnr(run, 18, bold=True, color=DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ЖКХ.Онлайн')
tnr(run, 22, bold=True, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Платформа сбора и обработки жалоб на жилищно-коммунальное хозяйство')
tnr(run, 13, color=SLATE)

doc.add_paragraph()

tbl = doc.add_table(rows=4, cols=2)
tbl.style  = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
meta = [
    ('Автор',          'Пугин Максим Витальевич'),
    ('Роли в проекте', 'Frontend Developer  •  Backend Developer  •  Data Engineer'),
    ('Сектор',         'Жилищно-коммунальное хозяйство (ЖКХ)'),
    ('Дата',           datetime.date.today().strftime('%d.%m.%Y')),
]
for i, (k, v) in enumerate(meta):
    set_cell(tbl.rows[i].cells[0], k, bold=True, bg='1E3A5F',
             color=RGBColor(255,255,255), size=11)
    set_cell(tbl.rows[i].cells[1], v, size=11)
    tbl.rows[i].cells[0].width = Cm(5)
    tbl.rows[i].cells[1].width = Cm(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  СОДЕРЖАНИЕ
# ══════════════════════════════════════════════════════════════════════════
h1('СОДЕРЖАНИЕ')
toc = [
    ('1.',    'Общие сведения о проекте'),
    ('2.',    'Проблема и цель'),
    ('3.',    'Роли и разделение ответственности'),
    ('4.',    'Архитектура системы'),
    ('5.',    'Описание страниц и функциональных модулей'),
    ('  5.1', 'Страница регистрации и входа'),
    ('  5.2', 'Личный кабинет пользователя'),
    ('  5.3', 'Трекинг жалобы по номеру тикета'),
    ('  5.4', 'Панель администратора — список жалоб'),
    ('  5.5', 'Панель администратора — статистика'),
    ('  5.6', 'Панель администратора — маршрут бригады'),
    ('  5.7', 'Панель администратора — граф инфраструктуры'),
    ('6.',    'Инновационные компоненты бэкенда'),
    ('  6.1', 'Автоматическая эскалация жалоб (asyncio)'),
    ('  6.2', 'Граф коммунальной инфраструктуры (алгоритмы)'),
    ('  6.3', 'Оптимизатор маршрута бригады (TSP)'),
    ('7.',    'Паттерны проектирования'),
    ('8.',    'Технологический стек'),
    ('9.',    'База данных и модели данных'),
    ('10.',   'REST API — ключевые эндпоинты'),
    ('11.',   'Жизненный цикл разработки'),
    ('12.',   'Стандарты и нормативная база'),
    ('13.',   'Среды разработки'),
    ('14.',   'Развёртывание (деплой)'),
    ('15.',   'Ключевые показатели эффективности (KPI)'),
    ('16.',   'Заключение'),
]
for num, title in toc:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title}')
    tnr(run, 11, bold=(not num.startswith('  ')))
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  1. ОБЩИЕ СВЕДЕНИЯ
# ══════════════════════════════════════════════════════════════════════════
h1('1. Общие сведения о проекте')

tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Table Grid'
rows_g = [
    ('Наименование',      'ЖКХ.Онлайн — Платформа сбора и обработки жалоб на ЖКХ'),
    ('Класс проекта',     'Веб-платформа (SPA-фронтенд + REST API бэкенд + БД + фоновые сервисы)'),
    ('Тип системы',       'Полностековое веб-приложение (Full-Stack Web Application)'),
    ('Целевая аудитория', 'Жители многоквартирных домов, операторы и администраторы УК'),
    ('Репозиторий',       'GitHub — приватный репозиторий, ветка main'),
]
for i, (k, v) in enumerate(rows_g):
    set_cell(tbl.rows[i].cells[0], k, bold=True, bg='DBEAFE', size=10)
    set_cell(tbl.rows[i].cells[1], v, size=10)

# ══════════════════════════════════════════════════════════════════════════
#  2. ПРОБЛЕМА И ЦЕЛЬ
# ══════════════════════════════════════════════════════════════════════════
h1('2. Проблема и цель')

h2('2.1. Проблема')
para(
    'В сфере жилищно-коммунального хозяйства граждане, столкнувшиеся с аварийными '
    'ситуациями или неисправностями — протечки кровли, отключение отопления, '
    'поломки лифтов, перебои в электроснабжении, — вынуждены лично обращаться '
    'в управляющую компанию, звонить по телефону или оформлять бумажные заявления. '
    'Данная схема имеет системные недостатки:'
)
bullet('Отсутствует единая база обращений — каждое отделение управляющей компании ведёт учёт самостоятельно.')
bullet('Нет возможности отследить статус поданной жалобы: гражданин не знает, принята ли заявка в работу.')
bullet('Операторы получают дублирующие звонки от разных жителей об одной и той же аварии.')
bullet('Просроченные заявки «теряются» — отсутствует механизм автоматического контроля сроков.')
bullet('Администрация не имеет аналитических инструментов для выявления проблемных зон и приоритизации инцидентов.')
bullet('Маршрут выезда аварийной бригады планируется вручную, без учёта оптимального порядка объезда.')

h2('2.2. Цель')
para(
    'Создать централизованную цифровую веб-платформу, обеспечивающую:'
)
bullet('Единую точку приёма жалоб от граждан через браузер без установки дополнительных приложений.')
bullet('Автоматическую регистрацию, нумерацию (уникальный тикет-номер) и полный трекинг каждого обращения.')
bullet('Личный кабинет гражданина с историей заявок, статусами и возможностью подачи с геолокацией.')
bullet('Рабочее место администратора с инструментами управления, аналитикой, картами и маршрутизацией.')
bullet('Автоматическую эскалацию просроченных жалоб без участия оператора на основе порогов по категории.')
bullet('Предиктивный анализ зон риска на основе графа коммунальной инфраструктуры с визуализацией.')

# ══════════════════════════════════════════════════════════════════════════
#  3. РОЛИ
# ══════════════════════════════════════════════════════════════════════════
h1('3. Роли и разделение ответственности')

para('Проект разработан одним человеком, совмещающим три профессиональные роли:', bold=True)
doc.add_paragraph()

tbl3 = doc.add_table(rows=4, cols=3)
tbl3.style = 'Table Grid'
table_hdr(tbl3, ['Роль', 'Зона ответственности', 'Реализованные компоненты'])
roles = [
    ('Frontend\nDeveloper',
     'Разработка пользовательского интерфейса, вёрстка, клиентская логика, UX-дизайн',
     'index.html (вся разметка SPA), style.css (адаптивная вёрстка), '
     'app.js (логика всех страниц и вкладок), api.js (HTTP-клиент). '
     'Интерактивные карты Leaflet. D3.js force-directed граф. Chart.js-дашборд. '
     'sessionStorage для восстановления вкладок. Кастомные CSS-тултипы.'),
    ('Backend\nDeveloper',
     'REST API, бизнес-логика, аутентификация, ORM-модели, фоновые сервисы, деплой',
     'FastAPI-роутеры: auth.py, complaints.py, admin.py, users.py. '
     'SQLAlchemy 2.0 async-модели (User, Complaint, Category, StatusHistory). '
     'JWT-аутентификация (python-jose + passlib). '
     'Паттерны: Facade, Observer, Command. '
     'Asyncio-автоэскалация. Docker Compose + Nginx.'),
    ('Data\nEngineer',
     'Алгоритмы обработки данных, графовая аналитика, пространственные вычисления',
     'infrastructure_graph.py: BFS, Дейкстра, итеративный Тарьян, '
     'PageRank, кластеризация связных компонент. '
     'route_optimizer.py: TSP ближайшего соседа, формула Хаверсина. '
     'Предиктивная оценка зон риска. Нормализация PageRank в [0,1].'),
]
for i, (r, z, c) in enumerate(roles, 1):
    set_cell(tbl3.rows[i].cells[0], r, bold=True, bg='EFF6FF', size=10)
    set_cell(tbl3.rows[i].cells[1], z, size=10)
    set_cell(tbl3.rows[i].cells[2], c, size=10)

# ══════════════════════════════════════════════════════════════════════════
#  4. АРХИТЕКТУРА
# ══════════════════════════════════════════════════════════════════════════
h1('4. Архитектура системы')

h2('4.1. Общая схема')
para(
    'Платформа построена по трёхзвенной клиент-серверной архитектуре. '
    'Статический фронтенд (Single Page Application на Vanilla JavaScript) '
    'взаимодействует с FastAPI-бэкендом через REST API по протоколу HTTP/JSON. '
    'Бэкенд хранит данные в PostgreSQL через асинхронный ORM (SQLAlchemy 2.0). '
    'В продакшне все три компонента изолированы в Docker-контейнерах и управляются '
    'через Docker Compose, взаимодействуя через внутреннюю Docker-сеть.'
)

h2('4.2. Слои приложения')
layers = [
    ('Представление\n(Frontend)',
     'Статические файлы (HTML/CSS/JS) раздаются Nginx. '
     'SPA реализован без фреймворков: модульная архитектура JS с разделением '
     'api.js (HTTP-клиент, все fetch-вызовы) и app.js (вся логика UI). '
     'Маршрутизация внутри SPA реализована через показ/скрытие div-секций '
     'без перезагрузки страницы. Авторизационный токен хранится в localStorage.'),
    ('API-слой\n(FastAPI)',
     'Четыре APIRouter: auth.py, complaints.py, admin.py, users.py. '
     'Middleware: CORS-политика, HTTPBearer для JWT. '
     'Зависимости (Depends): get_db возвращает AsyncSession, '
     'get_current_user извлекает пользователя из JWT, '
     'get_current_admin добавляет проверку is_admin.'),
    ('Бизнес-логика\n(Services + Patterns)',
     'ComplaintFacade — единая точка управления жалобами, скрывает ORM от роутеров. '
     'EventBus — асинхронная шина событий при смене статуса. '
     'Command-объекты — инкапсуляция бизнес-операций. '
     'InfrastructureGraph — весь графовый анализ. RouteOptimizer — TSP.'),
    ('Данные\n(SQLAlchemy + PostgreSQL)',
     'Модели: User, Complaint, Category, StatusHistory. '
     'Асинхронный движок (create_async_engine + AsyncSession). '
     'Автомиграция при старте (Base.metadata.create_all). '
     'Сиды: категории с emoji-иконками и индивидуальными порогами эскалации.'),
    ('Фоновые задачи\n(asyncio)',
     'При startup FastAPI создаётся asyncio.create_task с бесконечным циклом. '
     'Каждые 60 секунд: SELECT просроченных жалоб → UPDATE status → '
     'INSERT в StatusHistory → EventBus.publish("complaint.escalated"). '
     'Полностью неблокирующее: await asyncio.sleep не тормозит обработку запросов.'),
]
tbl4 = doc.add_table(rows=5, cols=2)
tbl4.style = 'Table Grid'
for i, (k, v) in enumerate(layers):
    set_cell(tbl4.rows[i].cells[0], k, bold=True, bg='F0F9FF', size=10)
    set_cell(tbl4.rows[i].cells[1], v, size=10)

# ══════════════════════════════════════════════════════════════════════════
#  5. СТРАНИЦЫ И МОДУЛИ
# ══════════════════════════════════════════════════════════════════════════
h1('5. Описание страниц и функциональных модулей')

divider()
h2('5.1. Страница регистрации и входа')
role_chip('Frontend Developer  •  Backend Developer')

para(
    'Первый экран платформы. Реализован в виде двух форм-вкладок на одной странице: '
    '«Вход» (по умолчанию) и «Регистрация». Переключение между формами происходит '
    'без перезагрузки через добавление CSS-класса .active и анимацию '
    'opacity + transform. Страница одинаково корректно отображается '
    'на мобильных устройствах и десктопе.'
)

h3('Что создано на Frontend')
bullet('Форма входа: поля email + пароль, кнопка «Войти», ссылка на регистрацию.')
bullet('Форма регистрации: имя, email, пароль, подтверждение пароля.')
bullet('Клиентская валидация: проверка совпадения паролей до отправки на сервер.')
bullet('Блок ошибок (.auth-error): «Неверный email или пароль», «Email уже занят» — '
       'без перезагрузки страницы.')
bullet('После успешного входа: JWT сохраняется в localStorage, '
       'автоматический редирект в кабинет (обычный или admin в зависимости от роли).')
bullet('При обновлении страницы: GET /auth/me по сохранённому токену '
       'восстанавливает сессию без повторного входа.')

h3('Что создано на Backend')
bullet('POST /auth/register — хеширование пароля (bcrypt/passlib), '
       'создание записи User, генерация JWT, возврат токена.')
bullet('POST /auth/login — поиск пользователя по email, '
       'verify_password(passlib), генерация JWT.')
bullet('GET /auth/me — декодирование JWT, возврат объекта пользователя.')
bullet('Пароль никогда не хранится и не передаётся в открытом виде.')

divider()
h2('5.2. Личный кабинет пользователя')
role_chip('Frontend Developer  •  Backend Developer')

para(
    'Основной рабочий экран для граждан. Отображается после входа под обычным '
    'аккаунтом. Реализован как однострановый интерфейс с тремя вкладками. '
    'Состояние активной вкладки сохраняется в sessionStorage — '
    'при обновлении страницы пользователь возвращается на ту же вкладку.'
)

h3('Вкладка 1 — «Мои жалобы»')
bullet('Список всех жалоб текущего пользователя, отсортированных по дате (новые сверху).')
bullet('Цветовая индикация статуса: Новое — серый, В работе — синий, '
       'Решено — зелёный, Эскалировано — красный.')
bullet('Карточка жалобы: номер тикета, заголовок, адрес, категория, дата подачи, статус-чип.')
bullet('Кнопка «Подробнее» — модальное окно с полным описанием и хронологией '
       'смены статусов (таблица: статус, дата, комментарий оператора).')
bullet('Пустое состояние: иллюстрация «Жалоб пока нет» + кнопка перехода к подаче.')

h3('Вкладка 2 — «Подать жалобу»')
bullet('Поле «Заголовок» — краткое описание проблемы.')
bullet('Поле «Описание» — подробное текстовое описание (textarea).')
bullet('Поле «Категория» — выпадающий список, данные из GET /categories.')
bullet('Поле «Адрес» — текстовый ввод с кнопкой «Определить моё местоположение».')
bullet('Геолокация: Geolocation API браузера → координаты → '
       'reverse-геокодинг через Nominatim API → подставляет адрес в поле.')
bullet('Ручной ввод адреса: прямой геокодинг через Nominatim → '
       'координаты lat/lng передаются вместе с жалобой в POST /complaints.')
bullet('После успешной подачи: модальное окно с номером тикета (ЖКХ-{год}-{номер}), '
       'жалоба мгновенно появляется в списке.')

h3('Вкладка 3 — «Мой профиль»')
bullet('Отображение: имя, email, дата регистрации.')
bullet('Статистика: общее число жалоб, количество решённых, '
       'вычисляемый процент решённых.')
bullet('Кнопка «Выйти» — очищает localStorage, перенаправляет на экран входа.')

h3('Backend-эндпоинты кабинета')
bullet('GET /complaints — список жалоб текущего пользователя (фильтр user_id из JWT).')
bullet('POST /complaints — создание жалобы: title, description, address, lat, lng, category_id.')
bullet('GET /complaints/{id} — детали + история статусов (JOIN с status_history).')
bullet('GET /categories — справочник категорий для выпадающего списка.')

divider()
h2('5.3. Трекинг жалобы по номеру тикета')
role_chip('Frontend Developer  •  Backend Developer')

para(
    'Публичная страница без авторизации. Позволяет любому гражданину отследить '
    'статус обращения, зная его уникальный номер тикета (например, ЖКХ-2024-00042). '
    'Реализована как отдельная секция на главной странице с полем ввода и кнопкой «Найти».'
)
bullet('Поле ввода принимает номер тикета в свободном формате.')
bullet('Результат — карточка: заголовок, категория, дата подачи, '
       'текущий статус с цветовым кодом, последний комментарий оператора.')
bullet('Если тикет не найден — блок с сообщением «Жалоба не найдена».')
bullet('Backend: GET /complaints/track/{ticket_number} — публичный эндпоинт '
       '(без Depends(get_current_user)), возвращает ограниченный набор полей '
       'без персональных данных заявителя.')

divider()
h2('5.4. Панель администратора — список жалоб')
role_chip('Frontend Developer  •  Backend Developer')

para(
    'Первая вкладка административного кабинета. Отображает все жалобы всех '
    'пользователей системы с расширенными возможностями управления. '
    'Доступна только пользователям с флагом is_admin=true.'
)

h3('Интерфейс и функциональность')
bullet('Таблица жалоб: тикет, пользователь, категория, адрес, дата, статус, столбец действий.')
bullet('Фильтрация по статусу и категории — выпадающие списки, фильтрация на клиенте.')
bullet('Live-поиск по тексту — фильтрация по заголовку и адресу при наборе символов.')
bullet('Смена статуса прямо в строке таблицы — PATCH-запрос без перехода на другую страницу.')
bullet('Ручная эскалация — кнопка с диалогом ввода причины (POST /admin/complaints/{id}/escalate).')
bullet('Удаление жалобы — кнопка с confirm-диалогом (DELETE /admin/complaints/{id}).')
bullet('Модальное окно — полная карточка жалобы с хронологией смены статусов.')

h3('Backend-эндпоинты')
bullet('GET /admin/complaints — все жалобы (skip, limit, status, category_id). Requires admin.')
bullet('PATCH /admin/complaints/{id} — статус или назначение исполнителя (data.assigned_to).')
bullet('POST /admin/complaints/{id}/escalate — запись в StatusHistory с reason.')
bullet('DELETE /admin/complaints/{id} — каскадное удаление с историей.')

divider()
h2('5.5. Панель администратора — статистика и дашборд')
role_chip('Frontend Developer  •  Data Engineer')

para(
    'Вторая вкладка кабинета администратора. Визуализирует агрегированные данные '
    'о работе платформы в виде KPI-карточек и интерактивных диаграмм на Chart.js. '
    'Данные загружаются при каждом переходе на вкладку.'
)

h3('KPI-карточки')
bullet('Всего жалоб — общее число обращений в системе.')
bullet('Обработано — жалобы в статусе «Решено» + процент от общего числа.')
bullet('Эскалировано — количество автоматически эскалированных + доля.')
bullet('Среднее время реакции — медиана часов от подачи до первой смены статуса.')

h3('Диаграммы')
bullet('Donut-chart — распределение жалоб по статусам (цвета соответствуют статус-чипам).')
bullet('Горизонтальная bar-chart — топ-7 категорий по количеству жалоб.')
bullet('Line-chart — динамика поступления жалоб по дням за последние 30 дней.')

h3('Backend')
bullet('GET /admin/stats — агрегированная статистика: группировка по статусам и категориям, '
       'временной ряд (GROUP BY DATE(created_at)). Все вычисления на стороне БД.')

divider()
h2('5.6. Панель администратора — маршрут выезда бригады')
role_chip('Frontend Developer  •  Data Engineer')

para(
    'Третья вкладка кабинета администратора. Инструмент оперативного планирования '
    'маршрута объезда активных инцидентов. Позволяет диспетчеру ввести начальную '
    'точку (адрес или координаты) и мгновенно получить оптимальный маршрут '
    'с визуализацией на карте и пошаговым списком.'
)

h3('Элементы управления')
bullet('Ввод начальной точки: адрес (прямой геокодинг через Nominatim) '
       'или ручной ввод координат lat, lng.')
bullet('Фильтр по статусам жалоб: чекбоксы «Новые», «В работе», «Эскалированные».')
bullet('Фильтр по категории — учитывается при запросе.')
bullet('Кнопка «Построить маршрут» — запрос к GET /admin/route.')

h3('Карта (Leaflet.js + CARTO)')
bullet('Светлый тайловый слой CARTO — читаем на светлом фоне интерфейса.')
bullet('Синий маркер «Старт» — начальная точка диспетчера.')
bullet('Цветные маркеры инцидентов: цвет соответствует статусу жалобы, '
       'иконка категории (emoji) внутри кружка через L.divIcon.')
bullet('Красная ломаная линия — маршрут в порядке оптимального объезда.')
bullet('Попап при клике на маркер: порядковый номер в маршруте, тикет, '
       'заголовок, адрес, статус.')
bullet('fitBounds — карта автоматически масштабируется под все точки маршрута.')

h3('Пошаговый список маршрута')
bullet('Нумерованный список точек в оптимальном порядке объезда.')
bullet('Для каждой точки: тикет, адрес, расстояние до следующей точки (км).')
bullet('Итоговое расстояние всего маршрута.')

h3('Алгоритм (Data Engineer)')
para(
    'Алгоритм «Жадного ближайшего соседа» (Greedy Nearest Neighbor, TSP-эвристика):'
)
bullet('Начало из заданной точки диспетчера (lat, lng).')
bullet('На каждом шаге: выбирается ближайшая непосещённая точка.')
bullet('Расстояние вычисляется по формуле Хаверсина (точная дуга большого круга, '
       'не евклидова плоскость).')
bullet('Сложность: O(n²), достаточно для n < 500 инцидентов.')
bullet('Каждой точке присваивается поле dist_to_next_km — дистанция до следующей.')

divider()
h2('5.7. Панель администратора — граф коммунальной инфраструктуры')
role_chip('Frontend Developer  •  Data Engineer  •  Backend Developer')

para(
    'Четвёртая (наиболее технически сложная) вкладка кабинета администратора. '
    'Аналитический инструмент, строящий граф взаимосвязей между инцидентами '
    'и прогнозирующий зоны риска с помощью пяти графовых алгоритмов. '
    'Результат визуализируется в двух режимах: D3.js-граф и Leaflet-карта.'
)

h3('Элементы управления')
bullet('Радиус близости: 300 м / 500 м / 1 км / 2 км — '
       'жалобы в пределах этого расстояния одной категории соединяются ребром.')
bullet('Фильтр статусов: мультиселект.')
bullet('Переключатель «D3-граф» / «На карте» — два режима визуализации.')

h3('Режим 1: D3.js Force-Directed Graph')
bullet('Узлы = жалобы; рёбра = пары жалоб одной категории, '
       'расстояние между которыми <= выбранный радиус.')
bullet('Размер узла: от 20 до 28px пропорционально степени вершины.')
bullet('Цвет обводки: соответствует статусу жалобы.')
bullet('Внутри узла: emoji-иконка категории (лифт, труба, свет, и др.).')
bullet('Под узлом: номер тикета.')
bullet('Мосты (критические рёбра): красная сплошная линия + подпись «мост».')
bullet('Обычные рёбра: серая пунктирная линия.')
bullet('Точки сочленения (критические узлы): пунктирное кольцо вокруг узла.')
bullet('Force simulation: forceLink(dist=90) + forceManyBody(strength=−60) + '
       'forceX/Y(strength=0.07) + forceCollide — узлы остаются в видимой области.')
bullet('Tooltip при наведении: заголовок, адрес, статус, категория, '
       'степень (число связей), PageRank-важность.')

h3('Режим 2: Leaflet-карта')
bullet('Узлы отображаются в реальных географических координатах (lat, lng).')
bullet('Кастомный L.divIcon: цветной кружок (цвет = статус) с emoji-иконкой категории.')
bullet('Рёбра: L.polyline (красная сплошная для мостов, серая пунктирная для остальных).')
bullet('Popup при клике: тикет, заголовок, статус, адрес, степень вершины, '
       'предупреждение если это точка сочленения.')

h3('KPI-блок (под графом)')
bullet('5 карточек: Жалобы в графе, Связей, Кластеров, Мостов, Точек сочленения.')

h3('Блок кластеров')
bullet('Каждый кластер: ID, число участников, оценка риска (%), '
       'список тикетов-чипов (кликабельны — открывают модал жалобы).')
bullet('Кнопка «Выделить» — подсвечивает узлы кластера на D3-графе на 2.2 секунды.')
bullet('Если в кластере есть точки сочленения — блок-предупреждение с перечислением.')

# ══════════════════════════════════════════════════════════════════════════
#  6. ИННОВАЦИОННЫЕ КОМПОНЕНТЫ
# ══════════════════════════════════════════════════════════════════════════
h1('6. Инновационные компоненты бэкенда')

h2('6.1. Автоматическая эскалация жалоб')
role_chip('Backend Developer  •  Data Engineer')

para(
    'Фоновый asyncio-процесс, запускаемый один раз при старте FastAPI-приложения '
    '(событие on_event("startup")). Выполняется в бесконечном цикле, '
    'каждые 60 секунд сканируя базу данных на предмет просроченных жалоб.'
)

bullet('Алгоритм: выборка жалоб с status IN ("new", "in_progress") и '
       'created_at < NOW() - INTERVAL escalation_hours часов.')
bullet('Порог эскалации индивидуален для каждой категории (поле escalation_hours в Category):')
bullet('Лифты, аварийные ситуации — 12 часов.', level=1)
bullet('Отопление, электроснабжение — 24 часа.', level=1)
bullet('Водоснабжение, канализация — 36 часов.', level=1)
bullet('Прочие категории — 72 часа.', level=1)
bullet('При срабатывании: UPDATE status = "escalated" + '
       'INSERT в status_history (changed_by=NULL, comment="Автоматическая эскалация").')
bullet('EventBus.publish("complaint.escalated") — подписчики обрабатывают уведомления.')
bullet('Полностью неблокирующее: await asyncio.sleep(60) не занимает поток.')

h2('6.2. Граф коммунальной инфраструктуры')
role_chip('Data Engineer')

para(
    'Реализован в backend/app/services/infrastructure_graph.py (~200 строк). '
    'Принимает список жалоб с координатами, строит взвешенный граф '
    'и последовательно запускает набор алгоритмов.'
)

h3('Построение графа (O(n²))')
bullet('Узлы: каждая жалоба с валидными lat и lng.')
bullet('Рёбра: две жалобы соединяются ребром если: одна категория '
       'И расстояние (Хаверсин) <= proximity_km.')
bullet('Вес ребра = расстояние в километрах (используется алгоритмом Дейкстры).')
bullet('Смежности хранятся в dict[int, list[int]] для BFS и Тарьяна, '
       'и dict[int, dict[int, float]] для Дейкстры.')

h3('Алгоритмы')
tbl_alg = doc.add_table(rows=6, cols=3)
tbl_alg.style = 'Table Grid'
table_hdr(tbl_alg, ['Алгоритм', 'Сложность', 'Что возвращает / применение'])
algs = [
    ('BFS (обход в ширину)',
     'O(V + E)',
     'bfs_zone(start_id, max_hops): '
     'все узлы в пределах N рёбер от стартового инцидента '
     'с расстоянием в хопах. Зона поражения при аварии.'),
    ('Дейкстра (Dijkstra)',
     'O((V+E) log V)\nheapq',
     'dijkstra(start_id): кратчайшие пути (в км) по весовым рёбрам '
     'от стартового узла до всех остальных. '
     'Оценка «близости» угрозы в predict_risk_zone.'),
    ('Тарьян — мосты\nи точки сочленения\n(итеративная реализация)',
     'O(V + E)',
     'find_bridges_and_articulations(): '
     'bridges — рёбра, разрывающие граф при удалении; '
     'articulations — узлы-«перешейки» сети. '
     'Итеративно: deque заменяет стек рекурсии.'),
    ('PageRank\n(адаптированный)',
     'O(V * I)\nI = 40 итераций',
     'pagerank(damping=0.85): рейтинг важности инцидента '
     'пропорционально его связности. '
     'Нормализован в [0, 1]. Выводится в тултипе как «Важность».'),
    ('Кластеризация\n(BFS-компоненты)',
     'O(V + E)',
     'connected_components(): разбивка на связные компоненты. '
     'Для каждого кластера: risk_score = '
     'PageRank_sum / size. Выводится как процент риска.'),
]
for i, (a, c, r) in enumerate(algs, 1):
    set_cell(tbl_alg.rows[i].cells[0], a, bold=True, size=10)
    set_cell(tbl_alg.rows[i].cells[1], c, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(tbl_alg.rows[i].cells[2], r, size=10)

doc.add_paragraph()
para(
    'Важно: алгоритм Тарьяна реализован итеративно (без рекурсии). '
    'Причина: рекурсивная реализация падает с RecursionError при числе '
    'узлов более ~1000, что реально для крупного города. '
    'Стек рекурсии явно эмулируется через collections.deque.'
)

h2('6.3. Оптимизатор маршрута бригады (TSP)')
role_chip('Data Engineer')

para(
    'Реализован в backend/app/services/route_optimizer.py. '
    'Решает задачу коммивояжёра (Travelling Salesman Problem) '
    'для набора активных инцидентов методом эвристики ближайшего соседа.'
)
bullet('Алгоритм: Greedy Nearest Neighbor — на каждом шаге выбирается '
       'ближайшая ещё не посещённая точка.')
bullet('Метрика расстояния: формула Хаверсина (Haversine formula) — '
       'точная дуга большого круга, учитывающая кривизну Земли.')
bullet('Начальная точка: координаты диспетчера (задаются как параметр запроса).')
bullet('Результат: упорядоченный список инцидентов + суммарная длина маршрута в км.')
bullet('Каждой точке добавляется поле dist_to_next_km — расстояние до следующей.')
bullet('Сложность O(n²), достаточна для n < 500.')

# ══════════════════════════════════════════════════════════════════════════
#  7. ПАТТЕРНЫ
# ══════════════════════════════════════════════════════════════════════════
h1('7. Паттерны проектирования')
role_chip('Backend Developer')

tbl_p = doc.add_table(rows=4, cols=4)
tbl_p.style = 'Table Grid'
table_hdr(tbl_p, ['Паттерн', 'Файл', 'Ключевые классы / методы', 'Зачем применён'])
pats = [
    ('Facade\n(Фасад)',
     'patterns/\nfacade.py',
     'ComplaintFacade:\n'
     '• create_complaint(data, user_id)\n'
     '• change_status(id, status, admin, comment)\n'
     '• assign(id, user_id, admin_id)\n'
     '• escalate(id, reason)\n'
     '• get_complaint(id)\n'
     '• get_all_complaints(filters...)',
     'API-роутеры не работают с AsyncSession напрямую — '
     'только через фасад. Скрывает ORM-сложность. '
     'Упрощает замену ORM и тестирование.'),
    ('Observer\n(Наблюдатель)',
     'patterns/\nobserver.py',
     'EventBus:\n'
     '• subscribe(event_name, handler)\n'
     '• publish(event_name, data)\n\n'
     'События:\n'
     '• complaint.created\n'
     '• complaint.status_changed\n'
     '• complaint.escalated',
     'Декаплинг бизнес-логики от побочных эффектов. '
     'Смена статуса публикует событие; подписчики '
     '(рассылки, логи) добавляются независимо. '
     'Вся цепочка асинхронна.'),
    ('Command\n(Команда)',
     'patterns/\ncommands.py',
     'SubmitComplaintCommand\n'
     'ChangeStatusCommand\n'
     'AssignComplaintCommand\n'
     '→ метод execute()\n'
     '→ возвращает result',
     'Инкапсуляция операций как объектов. '
     'Единый формат для логирования. '
     'Основа для очереди задач и Undo-механизма.'),
]
for i, row in enumerate(pats, 1):
    for j, cell_txt in enumerate(row):
        set_cell(tbl_p.rows[i].cells[j],
                 cell_txt,
                 bold=(j == 0),
                 bg=('F0FDF4' if j == 0 else None),
                 size=9)

# ══════════════════════════════════════════════════════════════════════════
#  8. СТЕК
# ══════════════════════════════════════════════════════════════════════════
h1('8. Технологический стек')

tbl_s = doc.add_table(rows=11, cols=3)
tbl_s.style = 'Table Grid'
table_hdr(tbl_s, ['Слой / Категория', 'Технология', 'Назначение'])
stack = [
    ('Backend-фреймворк',  'FastAPI 0.110+',
     'ASGI-фреймворк, автогенерация OpenAPI-схемы, система зависимостей Depends'),
    ('ORM',                'SQLAlchemy 2.0 (async)',
     'Декларативные модели, AsyncSession, create_async_engine'),
    ('СУБД (продакшн)',    'PostgreSQL 15',
     'Основное реляционное хранилище, ACID-транзакции'),
    ('СУБД (разработка)',  'SQLite + aiosqlite',
     'Локальный режим без Docker: файловая БД'),
    ('Аутентификация',     'python-jose + passlib + bcrypt',
     'JWT (HS256, TTL 7 дней) + хеширование паролей'),
    ('Frontend',           'HTML5 + CSS3 + Vanilla JS (ES2022)',
     'SPA без фреймворков, модульная архитектура api.js / app.js'),
    ('Карты',              'Leaflet.js 1.9 + CARTO tiles',
     'Интерактивные карты маршрута бригады и графа инфраструктуры'),
    ('Графы (виз.)',       'D3.js v7',
     'Force-directed визуализация графа со стабилизацией узлов'),
    ('Диаграммы',          'Chart.js',
     'Дашборд администратора: donut, bar, line charts'),
    ('Контейнеризация',    'Docker + Docker Compose',
     'Изолированный деплой трёх сервисов: postgres, backend, frontend'),
]
for i, (l, t, n) in enumerate(stack, 1):
    set_cell(tbl_s.rows[i].cells[0], l, bold=True, bg='F8FAFC', size=10)
    set_cell(tbl_s.rows[i].cells[1], t, size=10)
    set_cell(tbl_s.rows[i].cells[2], n, size=10)

# ══════════════════════════════════════════════════════════════════════════
#  9. БАЗА ДАННЫХ
# ══════════════════════════════════════════════════════════════════════════
h1('9. База данных и модели данных')
role_chip('Backend Developer  •  Data Engineer')

tbl_db = doc.add_table(rows=5, cols=3)
tbl_db.style = 'Table Grid'
table_hdr(tbl_db, ['Таблица', 'Ключевые поля', 'Назначение'])
models = [
    ('users',
     'id, email (UNIQUE), password_hash,\nfull_name, is_admin (bool), created_at',
     'Аккаунты граждан и администраторов. '
     'is_admin=true открывает доступ к /admin/* эндпоинтам. '
     'password_hash — bcrypt, соль случайная.'),
    ('categories',
     'id, name, icon (emoji str),\nescalation_hours (int), description',
     'Справочник категорий жалоб. '
     'icon — emoji-иконка категории (лифт, труба, лампочка...). '
     'escalation_hours — индивидуальный порог автоэскалации.'),
    ('complaints',
     'id, ticket_number (UNIQUE), title,\ndescription, address, lat (float),\n'
     'lng (float), status, user_id FK,\ncategory_id FK, assigned_to FK, created_at',
     'Основная таблица обращений. '
     'ticket_number = "ЖКХ-{year}-{id:05d}". '
     'lat/lng хранятся для графового анализа и карт. '
     'status: new / in_progress / resolved / escalated.'),
    ('status_history',
     'id, complaint_id FK, old_status,\nnew_status, comment, changed_by FK,\nchanged_at',
     'Полная хронология смены статусов. '
     'changed_by=NULL при автоэскалации (системное действие). '
     'comment — причина смены или комментарий оператора.'),
]
for i, (t, f, d) in enumerate(models, 1):
    set_cell(tbl_db.rows[i].cells[0], t, bold=True, bg='FFFBEB', size=10)
    set_cell(tbl_db.rows[i].cells[1], f, size=9)
    set_cell(tbl_db.rows[i].cells[2], d, size=10)

# ══════════════════════════════════════════════════════════════════════════
#  10. REST API
# ══════════════════════════════════════════════════════════════════════════
h1('10. REST API — ключевые эндпоинты')
role_chip('Backend Developer')

tbl_api = doc.add_table(rows=18, cols=3)
tbl_api.style = 'Table Grid'
table_hdr(tbl_api, ['Метод + маршрут', 'Доступ', 'Описание'])
endpoints = [
    ('POST /auth/register',            'Публичный', 'Регистрация пользователя, возврат JWT-токена'),
    ('POST /auth/login',               'Публичный', 'Вход по email+пароль, возврат JWT'),
    ('GET /auth/me',                   'User/Admin','Профиль текущего пользователя по токену'),
    ('GET /categories',                'Публичный', 'Список категорий жалоб для выпадающих списков'),
    ('GET /complaints/track/{ticket}', 'Публичный', 'Трекинг по номеру тикета (без авторизации)'),
    ('GET /complaints',                'User',       'Жалобы текущего пользователя'),
    ('POST /complaints',               'User',       'Подача новой жалобы (с lat/lng)'),
    ('GET /complaints/{id}',           'User',       'Детали жалобы + история статусов'),
    ('GET /admin/complaints',          'Admin',      'Все жалобы с фильтрами (status, category_id, пагинация)'),
    ('PATCH /admin/complaints/{id}',   'Admin',      'Смена статуса или назначение исполнителя'),
    ('POST /admin/complaints/{id}/escalate','Admin', 'Ручная эскалация с причиной'),
    ('DELETE /admin/complaints/{id}',  'Admin',      'Удаление жалобы'),
    ('GET /admin/users',               'Admin',      'Список всех пользователей системы'),
    ('GET /admin/stats',               'Admin',      'Агрегированная статистика для дашборда'),
    ('GET /admin/route',               'Admin',      'TSP-маршрут бригады (lat, lng, statuses, category_id)'),
    ('GET /admin/graph/analysis',      'Admin',      'Полный графовый анализ (BFS/Dijkstra/Tarjan/PageRank/кластеры)'),
    ('GET /admin/graph/predict/{id}',  'Admin',      'Прогноз зоны риска для конкретного инцидента'),
]
for i, (m, a, d) in enumerate(endpoints, 1):
    bg_a = ('FEF2F2' if a == 'Admin' else ('EFF6FF' if a == 'User' else 'F0FDF4'))
    set_cell(tbl_api.rows[i].cells[0], m, bold=True, size=9)
    set_cell(tbl_api.rows[i].cells[1], a, size=9, bg=bg_a,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(tbl_api.rows[i].cells[2], d, size=9)

# ══════════════════════════════════════════════════════════════════════════
#  11. ЖИЗНЕННЫЙ ЦИКЛ
# ══════════════════════════════════════════════════════════════════════════
h1('11. Модель жизненного цикла разработки')

para(
    'Применяется спиральная модель жизненного цикла (Spiral Life Cycle Model, '
    'Б. У. Боэм, 1986). Каждый виток включает четыре квадранта: '
    'определение целей → оценка рисков → разработка и тестирование → '
    'планирование следующего витка. Модель выбрана из-за эволюционного характера '
    'требований: функциональность платформы наращивалась итеративно, '
    'без деградации ранее реализованных компонентов.'
)

tbl_lc = doc.add_table(rows=6, cols=4)
tbl_lc.style = 'Table Grid'
table_hdr(tbl_lc, ['Виток', 'Реализованная функциональность', 'Роли', 'Результат'])
spirals = [
    ('Виток 1\nMVP',
     'Регистрация и вход (JWT, bcrypt). '
     'Подача жалобы с геолокацией (Nominatim). '
     'Личный кабинет с историей. '
     'ORM-модели, сиды категорий.',
     'Backend\nFrontend',
     'Работающий MVP: граждане могут подавать жалобы'),
    ('Виток 2\nУправление',
     'Кабинет администратора: список, фильтры, '
     'смена статуса, удаление, ручная эскалация. '
     'Паттерны Facade, Observer, Command. '
     'Автоэскалация asyncio.',
     'Backend\nFrontend',
     'Операторы могут управлять жалобами'),
    ('Виток 3\nАналитика',
     'Дашборд с Chart.js: donut, bar, line. '
     'Трекинг по тикету без авторизации. '
     'GET /admin/stats.',
     'Data\nFrontend',
     'Видна статистика работы платформы'),
    ('Виток 4\nКарты + Маршрут',
     'Карта инцидентов на Leaflet. '
     'TSP-маршрутизатор выезда бригады. '
     'Ввод стартовой точки адресом/координатами.',
     'Data\nFrontend',
     'Бригада получает оптимальный маршрут'),
    ('Виток 5\nГраф',
     'InfrastructureGraph: BFS, Дейкстра, '
     'итеративный Тарьян, PageRank, кластеризация. '
     'D3.js force-graph + Leaflet-граф. '
     'Прогноз зон риска. KPI + кластеры.',
     'Data\nBackend\nFrontend',
     'Предиктивная аналитика зон риска'),
]
for i, (v, f, r, res) in enumerate(spirals, 1):
    set_cell(tbl_lc.rows[i].cells[0], v, bold=True, bg='EFF6FF', size=10)
    set_cell(tbl_lc.rows[i].cells[1], f, size=10)
    set_cell(tbl_lc.rows[i].cells[2], r, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(tbl_lc.rows[i].cells[3], res, size=10)

# ══════════════════════════════════════════════════════════════════════════
#  12. СТАНДАРТЫ
# ══════════════════════════════════════════════════════════════════════════
h1('12. Стандарты и нормативная база')

tbl_st = doc.add_table(rows=4, cols=3)
tbl_st.style = 'Table Grid'
table_hdr(tbl_st, ['Документ', 'Суть стандарта', 'Применение в проекте'])
stds = [
    ('Федеральный закон № 59-ФЗ\n«О порядке рассмотрения\nобращений граждан\nРоссийской Федерации»',
     'Устанавливает обязательный срок рассмотрения '
     'обращений — 30 дней. Регламентирует обязательность '
     'ответа и недопустимость потери заявки.',
     'Логика автоэскалации реализует дух ФЗ-59: '
     'жалоба не может «зависнуть» бесконечно. '
     'Система принудительно переводит её в статус '
     '«Эскалировано» при превышении срока, '
     'уведомляя ответственных.'),
    ('ГОСТ 19.102-77\n«Стадии разработки\nпрограмм и\nпрограммной документации»',
     'Регламентирует этапы ЖЦ ПО: техническое задание '
     '→ эскизный проект → технический проект '
     '→ рабочая документация → внедрение.',
     'Этапы проекта соответствуют витками спирали: '
     'виток 1 (ТЗ+ЭП), витки 2-3 (ТП), '
     'витки 4-5 (РД и внедрение). '
     'Docker Compose — среда внедрения.'),
    ('ГОСТ Р 52436-2005\n«Программная инженерия.\nРуководство по применению\nстандартов»',
     'Адаптированные требования для малых организаций: '
     'управление требованиями, конфигурациями, '
     'тестирование, документирование.',
     'Версионирование через Git (ветка main). '
     'Docker-образы как артефакты конфигурации. '
     'Ручное регрессионное тестирование '
     'на каждом витке спирали.'),
]
for i, (d, s, a) in enumerate(stds, 1):
    set_cell(tbl_st.rows[i].cells[0], d, bold=True, bg='F1F5F9', size=9)
    set_cell(tbl_st.rows[i].cells[1], s, size=9)
    set_cell(tbl_st.rows[i].cells[2], a, size=9)

# ══════════════════════════════════════════════════════════════════════════
#  13. СРЕДЫ РАЗРАБОТКИ
# ══════════════════════════════════════════════════════════════════════════
h1('13. Среды разработки')

tbl_ide = doc.add_table(rows=3, cols=3)
tbl_ide.style = 'Table Grid'
table_hdr(tbl_ide, ['IDE', 'Область применения', 'Ключевые возможности'])
ides = [
    ('WebStorm\n(JetBrains)',
     'Фронтенд:\nHTML, CSS, JavaScript,\nLeaflet.js, D3.js, Chart.js',
     'Автодополнение JS/CSS/HTML. Встроенный браузерный отладчик. '
     'ESLint-интеграция. Живой предпросмотр в браузере. '
     'Инспекция DOM. Работа с CDN-библиотеками. '
     'Профилировщик производительности.'),
    ('IntelliJ IDEA\n(JetBrains)',
     'Бэкенд:\nPython, FastAPI,\nSQLAlchemy, Docker',
     'Python-плагин с интроспекцией типов (PEP 484 / type hints). '
     'Интеграция с venv, pip, Docker Compose. '
     'HTTP-клиент для тестирования REST-эндпоинтов. '
     'Отладчик asyncio-кода. Database-инструмент для SQLite.'),
]
for i, (ide, a, f) in enumerate(ides, 1):
    set_cell(tbl_ide.rows[i].cells[0], ide, bold=True, bg='F0F9FF', size=10)
    set_cell(tbl_ide.rows[i].cells[1], a, size=10)
    set_cell(tbl_ide.rows[i].cells[2], f, size=10)

# ══════════════════════════════════════════════════════════════════════════
#  14. ДЕПЛОЙ
# ══════════════════════════════════════════════════════════════════════════
h1('14. Развёртывание (деплой)')
role_chip('Backend Developer')

h2('14.1. Продакшн — Docker Compose')
para(
    'Производственная конфигурация описывает три изолированных Docker-контейнера, '
    'взаимодействующих через общую внутреннюю сеть. '
    'Переменные окружения (пароли, секреты) вынесены в .env-файл, '
    'не попадающий в Git.'
)

tbl_dc = doc.add_table(rows=4, cols=3)
tbl_dc.style = 'Table Grid'
table_hdr(tbl_dc, ['Сервис (контейнер)', 'Образ / сборка', 'Конфигурация'])
services = [
    ('postgres',
     'postgres:15-alpine\n(официальный образ)',
     'Порт 5432 (внутренний, не проброшен наружу). '
     'Health-check: pg_isready каждые 5 сек. '
     'Named volume для персистентного хранения данных. '
     'Env: POSTGRES_USER, POSTGRES_PASSWORD, POSTGRES_DB.'),
    ('backend',
     'Dockerfile:\nFROM python:3.11-slim\npip install -r requirements.txt\nCMD uvicorn',
     'Порт 8000 (внутренний). '
     'depends_on: postgres с condition: service_healthy. '
     'Env: DATABASE_URL (asyncpg://...), SECRET_KEY. '
     'При старте: auto-migrate (create_all) + seed категорий.'),
    ('frontend',
     'Dockerfile:\nFROM nginx:alpine\nCOPY frontend/ /usr/share/nginx/html\nCOPY nginx.conf',
     'Порт 80 (проброшен наружу). '
     'nginx.conf: статика frontend/, '
     'location /api/ → proxy_pass http://backend:8000/. '
     'gzip-сжатие, кеширование статики.'),
]
for i, (s, img, p) in enumerate(services, 1):
    set_cell(tbl_dc.rows[i].cells[0], s, bold=True, bg='F0FDF4', size=10)
    set_cell(tbl_dc.rows[i].cells[1], img, size=9)
    set_cell(tbl_dc.rows[i].cells[2], p, size=9)

h2('14.2. Локальная разработка (без Docker)')
bullet('Точка входа: local_main.py — FastAPI монтирует frontend/ как StaticFiles.')
bullet('База данных: SQLite через aiosqlite (файл local.db в корне проекта).')
bullet('Переключение БД: переменная окружения USE_SQLITE=true.')
bullet('Hot-reload: uvicorn --reload через watchfiles (перезапуск при изменении .py).')
bullet('Команда запуска: python local_main.py (в активированном venv).')
bullet('Адрес: http://localhost:8000 — единая точка: API + статика.')

# ══════════════════════════════════════════════════════════════════════════
#  15. KPI
# ══════════════════════════════════════════════════════════════════════════
h1('15. Ключевые показатели эффективности (KPI)')

tbl_k = doc.add_table(rows=4, cols=4)
tbl_k.style = 'Table Grid'
table_hdr(tbl_k, ['KPI', 'Описание', 'Метод измерения', 'Целевое значение'])
kpis = [
    ('Доля обработанных\nжалоб без потери',
     'Процент жалоб, получивших хотя бы один '
     'статусный переход (вышедших из статуса «Новое»). '
     'Характеризует полноту обработки обращений.',
     'resolved_count / total_count × 100%\n\n'
     'Отображается в дашборде администратора '
     'в секции статистики.',
     '≥ 95 %'),
    ('Количество жалоб,\nобработанных\nавтоматически',
     'Число жалоб, переведённых фоновым процессом '
     'автоэскалации в статус «Эскалировано» '
     'без действий оператора. Показывает '
     'эффективность автоматизации.',
     'COUNT(*) FROM status_history\n'
     'WHERE changed_by IS NULL\n'
     'AND new_status = "escalated"\n\n'
     'Выводится в KPI-карточке дашборда.',
     'Фиксируется,\nцелевое — рост\nотносительно\nпрошлого периода'),
    ('Доступность\nпортала (Uptime)',
     'Процент времени, когда API-сервер доступен '
     'и отвечает без ошибок 5xx. '
     'Определяет надёжность платформы '
     'для конечных пользователей.',
     'GET /api/health → HTTP 200\n\n'
     'Мониторинг через Docker healthcheck '
     'каждые 30 секунд.',
     '≥ 99 %\nв рабочее время'),
]
for i, (k, d, m, t) in enumerate(kpis, 1):
    set_cell(tbl_k.rows[i].cells[0], k, bold=True, bg='DBEAFE', size=10)
    set_cell(tbl_k.rows[i].cells[1], d, size=10)
    set_cell(tbl_k.rows[i].cells[2], m, size=9)
    set_cell(tbl_k.rows[i].cells[3], t, bold=True, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER)

# ══════════════════════════════════════════════════════════════════════════
#  16. ЗАКЛЮЧЕНИЕ
# ══════════════════════════════════════════════════════════════════════════
h1('16. Заключение')

para(
    'В ходе проектной работы разработана и развёрнута полнофункциональная '
    'веб-платформа ЖКХ.Онлайн, решающая задачу централизованного приёма, '
    'учёта и контроля исполнения обращений граждан в сфере жилищно-коммунального '
    'хозяйства. Платформа обеспечивает полный жизненный цикл заявки: '
    'от онлайн-регистрации гражданином с геолокацией — до автоматической '
    'эскалации при нарушении сроков и предиктивного анализа зон риска.'
)
para(
    'Как Frontend Developer реализованы все экраны платформы: страница входа '
    'и регистрации, личный кабинет гражданина с тремя вкладками (жалобы, '
    'подача, профиль), публичный трекинг по тикету, административный кабинет '
    'с четырьмя инструментами (управление, дашборд, маршрут, граф). '
    'Использованы Leaflet.js, D3.js v7, Chart.js.'
)
para(
    'Как Backend Developer создан REST API на FastAPI: 17 эндпоинтов, '
    'JWT-аутентификация, SQLAlchemy 2.0 async ORM, три паттерна проектирования '
    '(Facade, Observer, Command), фоновая asyncio-задача автоэскалации, '
    'контейнеризация через Docker Compose.'
)
para(
    'Как Data Engineer реализованы пять графовых алгоритмов '
    '(BFS, Дейкстра, итеративный Тарьян, PageRank, кластеризация) '
    'и TSP-маршрутизатор с формулой Хаверсина. '
    'Итеративная реализация алгоритма Тарьяна решила проблему '
    'переполнения стека рекурсии Python при n > 1000 узлов.'
)
para(
    'Применение спиральной модели жизненного цикла позволило пройти '
    'пять итераций разработки, каждая из которых добавляла новый '
    'функциональный слой поверх устойчивого фундамента предыдущего. '
    'Проект задокументирован, контейнеризирован и готов к промышленному '
    'развёртыванию на любом сервере с Docker.'
)

# ══════════════════════════════════════════════════════════════════════════
#  Сохранение
# ══════════════════════════════════════════════════════════════════════════
out = 'Отчёт_ЖКХ_Онлайн_Пугин.docx'
doc.save(out)
print(f'Saved: {out}')
